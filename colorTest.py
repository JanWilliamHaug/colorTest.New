# This program will search for red colored tags to get the information it needs
#
# Instructions on how to use the program:
# 1. run program
# 2. click on button "Choose document" and choose a document
# 3. Do step 2 as many times as you want, to uadd as many documents as you like
# 4. After you are done choosing documents, click on the GenerateReport button
# 5. Then you can click on the "open generated report" button, which will automatically
# open up your word document report created from your documents
# 6. When you are done, click "End Program"


import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
from tkinter import *
from tkinter import filedialog
from typing import Tuple

import re

import copy
import time

#This libraries are for opening word document automatically
import os
import platform
import subprocess

import xlwings as xw

import pandas as pd

# reads the text in the document and use the getcoloredTXT function
def readtxt(filename, color: Tuple[int, int, int]):
    doc = docx.Document(filename)
    text10 = ""
    fullText = []
    new = []
    global everything
    everything = []  # list of tags and text

    for para in doc.paragraphs:
        # Getting the colored words from the doc
        if (getcoloredTxt(para.runs, color)):
            # Concatenating list of runs between the colored text to single a string
            sentence = "".join(r.text for r in para.runs)
            fullText.append(sentence)
            #print(sentence) # Prints everything in the terminal
            everything.append(sentence)
            text10 = sentence
            parent.append("".join(r.text for r in para.runs))

    #print(fullText)
    global hasChild # Will store the ones with a child tag
    global fullText2 # will store everything found
    global children
    # Finds the lines without a childTag
    filtered_L = [value for value in fullText if "[" not in value]
    filtered_L = [s.replace(": ", ":") for s in filtered_L]
    # Finds the lines with a childTag
    filtered_LCopy.extend(filtered_L)
    hasChild = [value for value in fullText if "[" in value]
    # will store everything found
    fullText2 = [value for value in fullText]
    fullText2 = [s.replace(": ", ":") for s in fullText2]
    fullText2Copy.extend(fullText2)



    return fullText, filtered_L, hasChild, filtered_LCopy, fullText2Copy, fullText2

def getcoloredTxt(runs, color): # Will look for colored text

    coloredWords, word = [], ""
    for run in runs:
        if run.font.color.rgb == RGBColor(*color):
            word += str(run.text) # Saves everything found


        elif word != "": # This will find the parentTags
            coloredWords.append(word)
            parentTags.append(word)
            parents.append(word)
            word = ""


    if word != "": # This will find the parentTags
        coloredWords.append(word + "\n")
        #word = removeAfter(word)
        child.append(word)
        withChild.append(word)



    return coloredWords #returns everything found


# def openFile(): #This will let the user pick a document from their own directory
#     global filepath
#     global filepath2
#     filepath = filedialog.askopenfilename(initialdir="/",
#                                           title="",
#                                           filetypes= (("word documents","*.docx"),
#                                                       ("all files","*.*")))
#     file = open(filepath,'r')
#     #print(filepath)
#     file.close()
#     # Will store the filepath to the document as a string
#     filepath2 = str(filepath)
#
#     return filepath2, filtered_L

def generateReport(): #Will generate the report for tags
    global filepath
    global filepath2
    filepath = filedialog.askopenfilename(initialdir="/",
                                          title="",
                                          filetypes= (("word documents","*.docx"),
                                                      ("all files","*.*")))
    file = open(filepath,'r')
    #print(filepath)
    file.close()
    # Will store the filepath to the document as a string
    filepath2 = str(filepath)

    # return filepath2, filtered_L

    fullText = readtxt(filename=filepath2,
                       color=(255, 0, 0))
    #filtered_L = readtxt(filename=filepath2, #For future use
    #                   color=(255, 0, 0))
    fullText10 = str(fullText)
    s = ''.join(fullText10)
    w = (s.replace (']', ']\n\n'))
    paragraph = report3.add_paragraph()
    filepath3 = str(filepath2.rsplit('/', 1)[-1]) # change filepath to something.docx
    filepath3 = filepath3.split('.', 1)[0] # removes .docx of the file name
    print(filepath3 + " added to the report")
    runner = paragraph.add_run("\n" + "Document Name: " + filepath3 + "\n")
    runner.bold = True #makes the header bold
    # w will be used in the future
    w = (w.replace ('([', ''))
    w = (w.replace (',', ''))
    w = (w.replace ('' '', ''))

    # creates a table
    table = report3.add_table(rows=1, cols=2)

    # Adds headers in the 1st row of the table
    row = table.rows[0].cells
    row[0].text = 'Child Tag'
    row[1].text = 'Parent Tag/tags'
    # Adding style to a table
    table.style = 'Colorful List'

    # Now save the document to a location
    report3.save('report3.docx')
    e = 0

    child2 = removeAfter(child) #removes everything after the parent tag if there is anything to remove
    # while loop until all the  parentTags has been added to the report


    parents2 = copy.deepcopy(parentTags) # copy of parent tags list
    parents2Copy.extend(parents2)
    childCopy = copy.deepcopy(child2)
    noParent = []
    noParent2 = []
    orphanChild = []
    orphanChildParent = []
    parents9000 = []

    parents2 = [s.replace(" ", "") for s in parents2] # gets rid of space
    while parentTags:
        row = table.add_row().cells # Adding a row and then adding data in it.
        row[0].text = parentTags[0] # Adds the parentTag to the table
        noParent.append(parentTags[0])


        if e < len(fullText2):  #as long as variable e is not higher than the lines in fullText2
            if fullText2[e] in filtered_LCopy: #filtered_L contains the parent tags without a child tag
                orphanChild.append(parentTags[0])
                parentTags.remove(parentTags[0]) # Removes that tag after use
                noParent2.append(" ")
                parents9000.append(" ")
                orphanChildParent.append(" ")
                row[1].text = " " # No parent tag, so adds empty string to that cell
                e += 1

            elif fullText2[e] not in filtered_LCopy:
                parentTags.remove(parentTags[0]) # Removes that tag after use
                if child2:
                    row[1].text = child2[0] #Adds childTag to table
                    e += 1
                    parents9000.append(child2[0])
                    noParent.append(child2[0])
                    child2.remove(child2[0])  # Removed that tag from the list
    """
    while parentTags: # In case there are any more parent tags left in the list
        row = table.add_row().cells # Adding a row and then adding data in it.
        row[0].text = parentTags[0]
        parentTags.remove(parentTags[0])
    while child2: #This is for orphan tags, but not finished
        row = table.add_row().cells # Adding a row and then adding data in it.
        row[1].text = child2[0]
        child2.remove(child2[0])
    """

    parents9.extend(parents9000)

    # Make sure everything is cleared before the program gets the next document
    child2.clear()
    parentTags.clear()
    child.clear()
    report3.save('report3.docx') #Saves in document "report3"

    global dicts11
    dicts11 = dict(zip(parents2, childCopy)) #creates a dictrionary if there is a child tag and parent tag
    dicts.update(dicts)

    noParent = [s.replace(" ", "") for s in noParent]
    #dicts3 = dict(zip(noParent, noParent2))


    orphanChild = [s.replace(" ", "") for s in orphanChild]

    dicts9000 = dict(zip(orphanChild, orphanChildParent)) # orphan dictionary
    orphanDicts.update(dicts9000)
    OrphanChild2.extend(orphanChild)


    #for x in parents2: # creates dicttionary for child tags and text
    #   text2 = removeParent(everything)  # child tag and text
    #text8 = [s.replace(" ", "") for s in text2]
    #  text3 = removechild(text2)  # only text list
    # text4 = removeText(text2)  # child tags
    ##text8 = [s.replace(" ", "") for s in text4]

    #dicts12 = dict(zip(parents2, text3))  # creates a dictionary with child tags and text
    #sorted(dicts3.keys())  # sorts the keys in the dictionary
    #dicts3.update(dicts12)



    #for x, y in dicts.items():
    #row = table.add_row().cells  # Adding a row and then adding data in it.
    #        row[0].text = x
    #       row[1].text = y
    # text1 = int(str(list(fullText)))
    # print(everything)
    text2 = removeParent(everything) # child tag and text
    #text2 = removechild(everything)  # parent tags and text

    # print(text2)
    #text3 = removeParent(text2)  # only text list
    #text9 = ('"""' + str(text2) + '"""')  # child tag and text
    text3 = removechild(text2)  # only text list
    # print(text3)
    text4 = removeText(text2) # child tags
    # print(text4) #only parent tag list
    #text7 = [s.replace(" ", "") for s in text3]
    text8 = [s.replace(" ", "") for s in text4]

    #dicts3 = dict(zip(parents2, childCopy))
    dicts3 = dict(zip(parents2, parents9000))

    dicts10.update(dicts3)

    dicts2 = dict(zip(parents2, text3)) # creates a dictionary with child tags and text
    dicts100 = copy.deepcopy(dicts2)
    sorted(dicts2.keys()) # sorts the keys in the dictionary
    dicts2Copy.update(dicts100)
    #print(dicts2)

    #row = table.add_row().cells  # Adding a row and then adding data in it.
    #row[0].text = ""
    #row[1].text = ""
    #print(dicts)
    #print(parents2)
    #print(fullText2[1])
    #print(filtered_L)
    #print(parents2)
    #print(dicts2)
    #print(dicts10)
    #print(dicts2Copy)
    #print(filtered_LCopy)
    #print(orphanDicts)
    #print(parents2)
    #print(text3)
    #print(dicts100)
    return filepath2, filtered_L
    return parents2, dicts2, dicts10, dicts2Copy, parents2Copy, fullText2, filtered_LCopy, dicts3, orphanDicts, OrphanChild2

def generateReport2():
    #print("parents2Copy")
    #print(parents2Copy)
    #print("dicts2Copy")
    #print(dicts2Copy)
    #print("dicts10")
    #print(dicts10)
    #print("filtered_LCopy")
    #print(filtered_LCopy)
    #print("fullText2Copy")
    #print(fullText2Copy)

    # declaring counters
    m = 0
    k = 0
    i = 0
    o = 0
    z = 0

    orphanTagText = removechild(filtered_LCopy)

    #print(dicts10)
    #print(dicts2Copy)
    while m < len(parents2Copy):
        #print(m)
        if fullText2Copy[k] not in filtered_LCopy:
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1
                for key, value in dicts2Copy.items():
                    #for key, value in dicts2Copy.items() and key, value in dicts3.items(): #work on this here and try
                    report3.add_paragraph("\n")
                    report3.add_paragraph(key)
                    stringKey = str(key)
                    stringKey2 = (stringKey.replace(' ', ''))
                    report3.add_paragraph(value)
                    m += 2
                    if k < len(fullText2Copy):
                        if fullText2Copy[k] not in filtered_LCopy:  # check if it is an orphan tag
                            k += 1
                            if str(stringKey) in dicts10:

                                text = dicts10[str(stringKey2)]
                                PTags = text.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                PTags.pop()

                                for x in PTags:
                                    #print(x)
                                    report3.add_paragraph(x, style='List Bullet')
                                    keyCheck = (x.replace('[', ''))
                                    keyCheck2 = (keyCheck.replace(']', ''))
                                    keyCheck3 = (keyCheck2.replace(']', ''))
                                    keyCheck4 = (keyCheck3.replace(' ', ''))
                                    #print(keyCheck4)
                                    if keyCheck4 in dicts2Copy:
                                        para = report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                                        para.paragraph_format.left_indent = Inches(0.25) # adds indentation ot text
                                    else:

                                        report3.add_paragraph("No text found", style='List Bullet')

                            else:
                                # m += 2
                                report3.add_paragraph("No parent tag found", style='List Bullet')
                                pass
                        else:
                            report3.add_paragraph(key + " is an orphan tag")
                            k += 1
                        #m += 2
                        #k += 1
        elif fullText2Copy[k] in filtered_LCopy:
            #print(orphanDicts)
            report3.add_paragraph("\n")
            report3.add_paragraph(parents2Copy[i])
            report3.add_paragraph(orphanTagText[o])
            o += 1
            report3.add_paragraph(parents2Copy[i] + " is an orphan tagg")
            m += 1
            k += 1
            i += 1
            del dicts2Copy[list(dicts2Copy.keys())[0]] # deletes the first item in dicts2Copy

    print("Report Generated")
    print("You can now open up your report")
    report3.save('report3.docx')
    return dicts2Copy

    """
        elif not dicts2Copy: # this is for orphan tags
            dict3 = dict(dicts2.items() - dicts3.items())
            for key, value in dicts3.items():
                report3.add_paragraph("\n")
                report3.add_paragraph(key)
                report3.add_paragraph(value)
                report3.add_paragraph(key + " is an orphan tags")
                m += 1
    """


def removeParent(text): #removes parent tags or child tags
    childAfter = []
    for line in text:
        childAfter = [i.rsplit('[', 1)[0] for i in text] # removes parent tags
        childAfter = [re.sub("[\(\[].*?[\)\]]", "", e) for e in childAfter]  # removes parent tags that are left
        childAfter = [re.sub("[\{\[].*?[\)\}]", "", e) for e in childAfter]  # removes "pass", "fail", etc.
    return childAfter


def removeText(text6): #this should remove everything before the parent tag
    childAfter = [s.split(None, 1)[0] for s in text6]
    return childAfter

def removeAfter(childtags): #removes everything after the  tag, example "pass"
    seperator = ']'
    childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]
    return childAfter

def removechild(text): #removes child, this one needs fixing
    mylst = []
    mylst = [s.split(None, 1)[1] for s in text]
    return mylst

# This function will open up the report automatically
def getDocument():
    if platform.system() == 'Darwin':
        subprocess.check_call(['open', 'report3.docx'])
    elif platform.system() == 'Windows':
        os.startfile('report3.docx')
    # os.startfile(report3) # try either one for windows if the first option gives error
    else:
        subprocess.call('xdg-open', report3)

# Creates an excel report
def createExcel():
    book_arr = xw.App().books
    wb = book_arr.add()
    #wb = xw.Book() # Creating an new excel file.
    # Select the first excel sheet, and rename it
    excelReport = wb.sheets["Sheet1"]

    report = "report"
    #excelReport.name = report
    excelReport.range("B1").value = "Report"
    excelReport.range("B1").api.Font.Size = 18 # Change font size
    excelReport.range("B1").api.Font.ColorIndex = 2 # Change font color
    excelReport.range('A1:S1').color = (0, 0, 255) # Change cell background color



    # creating a Dataframe object from a list
    # of tuples of key, value pair
    df = pd.DataFrame(list(dicts2Copy.items()))

    # Dictionary For child and parent tag
    df2 = pd.DataFrame(list(dicts10.items()))

    # For childTag -Text
    excelReport.range("A3").value = df

    # Adding childTag header
    excelReport.range("B3").value = 'Child Tag'
    excelReport.range("B3").api.Font.Size = 14 # Change font size
    excelReport.range("B3").api.Font.ColorIndex = 2 # Change font color
    excelReport.range('B3:B3').color = (255, 0, 0) # Change cell background color

    # Adding Text header
    excelReport.range("C3").value = 'Text'
    excelReport.range("C3").api.Font.Size = 14 # Change font size
    excelReport.range("C3").api.Font.ColorIndex = 2 # Change font color
    excelReport.range('C3:C3').color = (0,255,0) # Change cell background color

    # For the childTag - parentTag
    excelReport.range("D3").value = df2

    # Adding parentTag header
    excelReport.range("F3").value = 'Parent Tag'
    excelReport.range("F3").api.Font.Size = 14 # Change font size
    excelReport.range("F3").api.Font.ColorIndex = 2 # Change font color
    excelReport.range('F3:F3').color = (128, 128, 128) # Change cell background color

    # Adding childTag header
    excelReport.range("E3").value = 'Child Tag'
    excelReport.range("E3").api.Font.Size = 14 # Change font size
    excelReport.range("E3").api.Font.ColorIndex = 2 # Change font color
    excelReport.range('E3:E3').color = (255, 0, 0) # Change cell background color

    wb.sheets["Sheet1"].autofit()

    wb.save('report.xlsx') # Saving excel report as 'report.xlsx'

if __name__ == '__main__':
    # Creates a word document, saves it as "report 3, and also adds a heading
    report3 = Document()
    report3.add_heading('Report', 0) #create word document
    paragraph = report3.add_paragraph()
    report3.save('report3.docx')
    dicts2Copy = {} # This will hold the dicts2 content in all documents

    global parents2Copy # parents2 list copy
    parents2Copy = []

    global filtered_L # Will store the ones without a child tag
    filtered_L = []

    global filtered_LCopy
    filtered_LCopy = []

    global fullText2Copy
    fullText2Copy = []

    global parents2 #list of parent tags or child tags
    parents2 = []

    # creates a dict for parent and child tags
    global dicts
    dicts = {}

    global OrphanChild2
    OrphanChild2 = []


    global dicts10
    dicts10 = {}
    global dicts3
    dicts3 = {}  # will hold parentTag and text, Orphan tags
    global dicts2
    dicts2 = {}  # will hold parentTag and text
    global orphanDicts
    orphanDicts = {}  # orphan dictionary

    global parents9
    parents9 = []

    # declaring different lists that will be used to store, tags and sentences
    parentTags = []
    parent = []  # This will be used to store everything
    child = [] # Used to Store child tags
    noChild = []  # Used to Store parentTags with no child
    withChild = [] # Used to Store parentTags with child tag
    parents = [] #Will be used for future function

    global orphanTagText
    orphanTagText = []  # Will be used to hold text of orphanChildTags


    # Creates the gui
    window = Tk(className=' TARGEST')
    # set window size #
    window.geometry("250x135")
    # Creates button 1
    Button(window, text="Choose Document ", command=generateReport).pack()
    # Creates button 2
    Button(window, text="Generate Report ", command=generateReport2).pack()
    # Creates button 3
    getDoc = Button(window, text="Open Generated Report", command=getDocument)
    getDoc.pack()
    # Creates Excel button button 4
    button = Button(text="Create Excel Report", command=createExcel)
    button.pack()
    # Creates button 5
    button = Button(text="End Program", command=window.destroy)
    button.pack()


    window.mainloop()
